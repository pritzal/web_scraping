from bs4 import BeautifulSoup
import requests
from docx import Document


class ContentA:
    def data(self):
        print("data")


class Bbc(ContentA):
    def __init__(self, k, docu):
        self.url = "https://www.bbc.co.uk/search?q=" + k.replace(" ", '+') + "&page="
        self.docu = docu
        soup = BeautifulSoup(requests.get(self.url + '1').text, "html.parser")
        self.p_no = int(list(soup.ol.strings)[-1])
        docu.add_heading("NEWS", 1)

    def data(self):
        for i in range(1, self.p_no + 1):
            doc = BeautifulSoup(requests.get(self.url + str(i)).text, "html.parser")
            li = doc.find("ul", role="list", class_="ssrcss-1020bd1-Stack e1y4nx260")
            for j in li.contents:
                t_4 = j.find("div", spacing="4")
                p = list(t_4.strings)
                if (p[3] != "News"):
                    continue
                t_2 = j.find("div", spacing="2")
                link = t_2.a['href']  # link for news
                news = t_2.a.span.p.span.string
                news_2 = t_2.contents[1].string
                self.docu.add_heading(news, 2)
                self.docu.add_paragraph(news_2)
                self.docu.add_paragraph(link)
                last = str("Time =" + p[1] + "  " + "Section =" + p[-1])
                self.docu.add_paragraph(last)


class Wiki(ContentA):
    url = str()
    stp = str()

    def __init__(self, kw, docu):  # kw stands for keyword
        self.url = "https://en.wikipedia.org/wiki/" + kw.replace(" ", "_")
        self.doc = BeautifulSoup(requests.get(self.url).text, "html.parser")
        self.docu = docu
        docu.add_heading(kw, 0)
        docu.add_heading(self.url, 1)

    def data(self):
        self.q = self.doc.find_all(["p", "h3", "h2", "ul"])
        for i in self.q:
            if (i.name == "h2" and list(i.strings)[0] == 'See also'):
                break
            # if(i.string=="Contents"):
            # continue
            if (i.name == "p"):
                k = True
                for j in i.strings:
                    if (j[0] == '[' and j[-1] == ']'):
                        continue
                    elif (j == "["):
                        k = False
                        continue
                    elif (j == ']'):
                        k = True
                        continue
                    if (k == True):
                        self.stp = self.stp + str(j)
                self.docu.add_paragraph(self.stp)
                self.stp = ''
            elif (i.name == "h3"):
                k = True
                for j in i.strings:
                    if (j[0] == '[' and j[-1] == ']'):
                        continue
                    elif (j == "["):
                        k = False
                        continue
                    elif (j == ']'):
                        k = True
                        continue
                    if (k == True):
                        self.stp = self.stp + str(j)
                self.docu.add_heading(self.stp, 3)
                self.stp = ''
            elif (i.name == "h2"):
                k = True
                for j in i.strings:
                    if (j[0] == '[' and j[-1] == ']'):
                        continue
                    elif (j == "["):
                        k = False
                        continue
                    elif (j == ']'):
                        k = True
                        continue
                    if (k == True):
                        self.stp = self.stp + str(j)
                self.docu.add_heading(self.stp, 2)
                self.stp = ''
            elif (i.name == 'ul'):
                k = True
                for j in i.strings:
                    if (j[0] == '[' and j[-1] == ']'):
                        continue
                    elif (j == "["):
                        k = False
                        continue
                    elif (j == ']'):
                        k = True
                        continue
                    if (k == True):
                        self.stp = self.stp + str(j)
                self.docu.add_paragraph(self.stp)
                self.stp = ''
